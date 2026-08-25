"""
Обезличивание Word (.docx) с сохранением вёрстки.

Обходит ВСЕ текстовые места (тело, таблицы, колонтитулы), прогоняет каждый
параграф через ядро анонимизации и вписывает результат обратно. Метаданные и
текст правок-удалений вычищаются (см. metadata.py).

Ключевая деталь корректности — ПДн, разорванная по runs. Word режет текст
параграфа на runs по форматированию: «Иван **Петров**» может лежать в двух runs.
Поэтому анонимизируем ТЕКСТ ПАРАГРАФА ЦЕЛИКОМ, а не отдельные runs. Если ПДн
найдена, результат кладём в первый run, остальные очищаем: внутрипараграфное
форматирование в таких (и только таких) параграфах не сохраняется — это
осознанный безопасный компромисс, приватность важнее.
"""
import io
import logging

from docx import Document

from app.anonymizer import anonymize_text

from .metadata import detect_review_leftovers, scrub_docx_metadata

logger = logging.getLogger(__name__)


def _anonymize_paragraph(paragraph, disable_entities) -> int:
    """Обезличивает один параграф по его полному тексту. Число найденных сущностей."""
    if not paragraph.runs:
        return 0
    full = "".join(r.text for r in paragraph.runs)
    if not full.strip():
        return 0
    res = anonymize_text(full, disable_entities=disable_entities)
    if res["anonymized"] == full:
        return 0
    # ПДн найдена — весь обезличенный текст в первый run, остальные очищаем
    paragraph.runs[0].text = res["anonymized"]
    for r in paragraph.runs[1:]:
        r.text = ""
    return len(res["entities_found"])


def _mask_values_in_paragraph(paragraph, value_to_ph) -> int:
    """Заменяет в параграфе найденные значения ПДн на их плейсхолдеры.

    Используется для ячеек таблиц: детект идёт по всей строке (контекст строки),
    а маскирование — по значению внутри ячейки. Длинные значения первыми — чтобы
    одно значение не перезатирало часть другого."""
    if not paragraph.runs or not value_to_ph:
        return 0
    full = "".join(r.text for r in paragraph.runs)
    if not full.strip():
        return 0
    masked, n = full, 0
    for value in sorted(value_to_ph, key=len, reverse=True):
        if value and value in masked:
            masked = masked.replace(value, value_to_ph[value])
            n += 1
    if masked == full:
        return 0
    paragraph.runs[0].text = masked
    for r in paragraph.runs[1:]:
        r.text = ""
    return n


def _process_table_row(row, disable_entities, stats):
    """Обезличивает строку таблицы с ОБЩИМ контекстом строки.

    Ключевая правка утечки: «ИНН» в одной ячейке и значение в соседней. Детект
    идёт по объединённому тексту строки (ячейки через пробел — якорь видит ключ),
    затем найденные значения маскируются в параграфах ячеек."""
    cell_texts = [c.text for c in row.cells]
    combined = " ".join(t for t in cell_texts if t.strip())
    value_to_ph = {}
    if combined.strip():
        res = anonymize_text(combined, disable_entities=disable_entities)
        # placeholder->value  ->  value->placeholder (значения различны)
        value_to_ph = {v: k for k, v in res["mapping"].items()}
    for cell in row.cells:
        for para in cell.paragraphs:
            n = _mask_values_in_paragraph(para, value_to_ph)
            if n:
                stats["total"] += n
                stats["changed"] += 1
        for nested in cell.tables:          # вложенные таблицы
            for nrow in nested.rows:
                _process_table_row(nrow, disable_entities, stats)


def _process_container(container, disable_entities, stats):
    """Параграфы вне таблиц — по полному тексту; таблицы — по строкам (контекст)."""
    for para in container.paragraphs:
        n = _anonymize_paragraph(para, disable_entities)
        if n:
            stats["total"] += n
            stats["changed"] += 1
    for table in container.tables:
        for row in table.rows:
            _process_table_row(row, disable_entities, stats)


def anonymize_docx(data: bytes, disable_entities=None) -> tuple[bytes, dict]:
    """Обезличивает .docx. Возвращает (байты нового файла, сводка).

    Сводка — только счётчики и предупреждения, БЕЗ значений ПДн и mapping
    (де-анонимизирующие данные в вывод не попадают)."""
    document = Document(io.BytesIO(data))

    stats = {"total": 0, "changed": 0}
    _process_container(document, disable_entities, stats)
    for section in document.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer,
                   section.even_page_header, section.even_page_footer):
            _process_container(hf, disable_entities, stats)
    total_entities = stats["total"]
    paragraphs_changed = stats["changed"]

    warnings = detect_review_leftovers(document)
    scrub_docx_metadata(document)
    if warnings:
        logger.warning("docx: остаточные ревью-артефакты не обработаны в v1: %s", warnings)

    out = io.BytesIO()
    document.save(out)
    summary = {
        "format": "docx",
        "entities_found": total_entities,
        "paragraphs_changed": paragraphs_changed,
        "warnings": warnings,
    }
    return out.getvalue(), summary
