"""Генерирует тестовые документы с синтетическими ПДн в web-тест/тестовые_файлы/.

Запуск (из корня репозитория):
    python web-тест/сгенерировать_тестовые_файлы.py
"""
import os

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "тестовые_файлы")
os.makedirs(OUT, exist_ok=True)

# синтетические (вымышленные) ПДн
PERSON = "Сидоров Пётр Иванович"
FIELDS = [
    ("ФИО", PERSON),
    ("ИНН", "7707083893"),
    ("Паспорт", "45 08 731902"),
    ("СНИЛС", "112-233-445 95"),
    ("Телефон", "+7 (921) 555-18-43"),
    ("E-mail", "petr.sidorov@example.com"),
    ("Дата рождения", "15.03.1985"),
    ("Карта", "5536 9102 7741 3086"),
    ("Адрес", "г. Москва, ул. Ленина, дом 5, квартира 12"),
]
BODY = (
    f"Заявление. Я, {PERSON}, прошу оформить услугу. "
    "Мои данные: ИНН 7707083893, СНИЛС 112-233-445 95, паспорт серия 45 08 номер 731902. "
    "Телефон +7 (921) 555-18-43, почта petr.sidorov@example.com, дата рождения 15.03.1985. "
    "Карта для оплаты 5536 9102 7741 3086. "
    "Проживаю по адресу: Санкт-Петербург, Невский проспект, дом 18, квартира 47."
)


def make_docx():
    from docx import Document
    doc = Document()
    doc.add_heading("Тестовый документ с ПДн", level=1)
    doc.add_paragraph(BODY)
    doc.add_paragraph("Таблица данных (метка и значение в разных ячейках):")
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Поле"
    t.rows[0].cells[1].text = "Значение"
    for label, val in FIELDS:
        c = t.add_row().cells
        c[0].text = label
        c[1].text = val
    path = os.path.join(OUT, "тестовый_документ.docx")
    doc.save(path)
    return path


def make_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    if not os.path.exists(font_path):
        print("! нет DejaVuSans.ttf — PDF пропущен")
        return None
    pdfmetrics.registerFont(TTFont("DV", font_path))
    path = os.path.join(OUT, "тестовый_документ.pdf")
    c = canvas.Canvas(path, pagesize=A4)
    c.setFont("DV", 12)
    y = 800
    c.drawString(50, y, "Тестовый документ с ПДн")
    y -= 30
    c.setFont("DV", 11)
    for label, val in FIELDS:
        c.drawString(50, y, f"{label}: {val}")
        y -= 22
    c.save()
    return path


if __name__ == "__main__":
    try:
        print("docx:", make_docx())
    except Exception as e:  # noqa: BLE001
        print("docx пропущен:", e)
    try:
        print("pdf :", make_pdf())
    except Exception as e:  # noqa: BLE001
        print("pdf пропущен:", e)
