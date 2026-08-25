"""
Обработчик .docx: обезличивание тела/таблиц/колонтитулов + очистка метаданных.
Категория: integration. ПДн regex-овые (телефон/email/ИНН/дата) — модель не нужна.
"""
import io

import pytest
from docx import Document

from app.integrations.documents.docx_handler import anonymize_docx

from .conftest import SAMPLE, build_docx, extract_docx_text

pytestmark = pytest.mark.integration


class TestDocxLeaks:
    def test_no_pii_leaks_anywhere(self):
        """Ни одно значение ПДн не осталось в теле, таблице или колонтитуле."""
        out, summary = anonymize_docx(build_docx())
        text = extract_docx_text(out)
        for name, value in SAMPLE.items():
            assert value not in text, f"утечка {name} в .docx"
        assert summary["entities_found"] >= 4


class TestDocxTableContext:
    """Ключ ПДн в одной ячейке строки, значение — в соседней (реальная выгрузка).

    Раньше ячейки обезличивались изолированно, и паспорт без контрольной суммы
    в соседней ячейке утекал — контекста «Паспорт» рядом не было."""

    def test_row_context_masks_value_in_adjacent_cell(self):
        doc = Document()
        table = doc.add_table(rows=0, cols=2)
        rows = [("Паспорт", "45 08 731902"),   # без контрольной суммы — главный кейс
                ("СНИЛС", "112-233-445 95"),
                ("ИНН", "7707083893")]
        for label, val in rows:
            c = table.add_row().cells
            c[0].text = label
            c[1].text = val
        buf = io.BytesIO()
        doc.save(buf)

        out, _ = anonymize_docx(buf.getvalue())
        text = extract_docx_text(out)
        for label, val in rows:
            assert val.replace(" ", "") not in text.replace(" ", ""), \
                f"утечка {label} из соседней ячейки: {val}"

    def test_metadata_author_cleared(self):
        out, _ = anonymize_docx(build_docx(author="Секретный Автор"))
        d = Document(io.BytesIO(out))
        assert d.core_properties.author == ""

    def test_placeholders_present(self):
        out, _ = anonymize_docx(build_docx())
        text = extract_docx_text(out)
        assert "<PHONE>" in text
        assert "<EMAIL>" in text

    def test_negative_text_preserved(self):
        """Обычный текст (номер заказа, сумма) не должен затираться целиком."""
        out, _ = anonymize_docx(build_docx())
        text = extract_docx_text(out)
        assert "обычный текст" in text
        assert "Заказ" in text  # слово-метка на месте

    def test_labels_preserved(self):
        """Слова-подсказки рядом с ПДн сохраняются (читаемость для оператора)."""
        out, _ = anonymize_docx(build_docx())
        text = extract_docx_text(out)
        assert "Телефон" in text
        assert "ИНН" in text

    def test_output_is_valid_docx(self):
        out, _ = anonymize_docx(build_docx())
        # открывается без ошибок = валидный .docx
        Document(io.BytesIO(out))


class TestDocxPersonRequiresModel:
    @pytest.mark.requires_model
    def test_person_redacted(self):
        from docx import Document as Doc
        d = Doc()
        d.add_paragraph("Клиент Иван Петров обратился повторно.")
        buf = io.BytesIO()
        d.save(buf)
        out, _ = anonymize_docx(buf.getvalue())
        text = extract_docx_text(out)
        assert "Иван Петров" not in text
        assert "<PERSON>" in text
