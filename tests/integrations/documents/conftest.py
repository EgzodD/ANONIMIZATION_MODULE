"""
Фикстуры для тестов адаптера документов. Форсируют DOCUMENT_ENABLED=true и
пересобирают app.main, чтобы роут /anonymize/document был зарегистрирован.
Плюс билдеры фикстур-файлов (.docx через python-docx, .pdf через reportlab).

Тесты используют ПДн, которые ловятся regex-ом БЕЗ модели PERSON (телефон, email,
ИНН, дата) — так проверка утечек работает и в CI, где модели нет. Отдельные
проверки ФИО помечены requires_model.
"""
import importlib
import io

import pytest
from docx import Document
from fastapi.testclient import TestClient

import app.main as app_main_module
from app.config import settings


@pytest.fixture
def client():
    orig = settings.document_enabled
    settings.document_enabled = True
    try:
        importlib.reload(app_main_module)
        yield TestClient(app_main_module.app, raise_server_exceptions=False)
    finally:
        settings.document_enabled = orig
        importlib.reload(app_main_module)


# ── билдеры фикстур ─────────────────────────────────────────────────────────
# Синтетические ПДн (не реальные). Значения вынесены, чтобы тесты сверяли утечки.
SAMPLE = {
    "phone": "+7 999 123 45 67",
    "email": "p.sidorov@mail.ru",
    "inn": "7712345678",
    "date": "05.03.1990",
}
NEGATIVE = "Заказ №4509 на сумму 15000 рублей"


def build_docx(author="Иван Петров") -> bytes:
    """.docx с ПДн в теле, таблице и колонтитуле + автор в метаданных."""
    d = Document()
    d.core_properties.author = author
    d.add_paragraph(f"Телефон {SAMPLE['phone']}, родился {SAMPLE['date']}.")
    d.add_paragraph(NEGATIVE + " — обычный текст.")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "ИНН"
    t.rows[0].cells[1].text = f"ИНН {SAMPLE['inn']}"
    d.sections[0].footer.paragraphs[0].text = f"Почта {SAMPLE['email']}"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def build_pdf() -> bytes:
    """.pdf с ПДн (reportlab + DejaVuSans для кириллицы). importorskip если нет reportlab."""
    pytest.importorskip("reportlab")
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    import os
    if not os.path.exists(font_path):
        pytest.skip("нет шрифта DejaVuSans для кириллицы в PDF-фикстуре")
    pdfmetrics.registerFont(TTFont("DV", font_path))

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.setFont("DV", 12)
    y = 800
    for ln in [
        f"Телефон {SAMPLE['phone']}, ИНН {SAMPLE['inn']}.",
        f"Родился {SAMPLE['date']}, почта {SAMPLE['email']}.",
        NEGATIVE + " — обычный текст.",
    ]:
        c.drawString(50, y, ln)
        y -= 30
    c.save()
    return buf.getvalue()


def extract_docx_text(data: bytes) -> str:
    """Весь видимый текст .docx: тело, таблицы, колонтитулы."""
    d = Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    for s in d.sections:
        for p in s.footer.paragraphs:
            parts.append(p.text)
    return "\n".join(parts)
